from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import datetime

# ── Color Palette ─────────────────────────────────────────────────────────────
COLOR_ACCENT = RGBColor(0x00, 0x5B, 0xBB)   # Professional blue
COLOR_DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # Near-black
COLOR_LIGHT  = RGBColor(0x44, 0x44, 0x44)   # Dark gray for body
COLOR_MUTED  = RGBColor(0x77, 0x77, 0x77)   # Light gray for captions


def _set_heading_color(run, color: RGBColor):
    run.font.color.rgb = color


def _add_horizontal_rule(doc):
    """Add a thin horizontal line."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '005BBB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def _style_heading(heading, level=1, color=None):
    """Apply color to a heading paragraph."""
    for run in heading.runs:
        if color:
            run.font.color.rgb = color
        run.font.name = "Calibri"


def _add_chapter(doc, number, title, color=COLOR_ACCENT):
    """Add a numbered chapter heading."""
    h = doc.add_heading(f"{number}. {title}", level=1)
    _style_heading(h, 1, color)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    _add_horizontal_rule(doc)
    return h


def _add_body(doc, text, indent=False):
    """Add formatted body text."""
    if not text:
        return
    p = doc.add_paragraph(str(text))
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = COLOR_LIGHT
    return p


def _add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(str(text))
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = COLOR_LIGHT
    return p


def _add_key_value_table(doc, data: dict):
    """Render a dict as a styled two-column table."""
    if not data:
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Details"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = COLOR_ACCENT
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'E8F0FE')
        cell._tc.get_or_add_tcPr().append(shading)

    for key, val in data.items():
        row = table.add_row().cells
        row[0].text = str(key).capitalize()
        row[1].text = str(val)
        for r in row:
            for run in r.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"
    doc.add_paragraph()


def _add_viva_table(doc, viva_questions):
    """Render viva Q&A as a styled table."""
    if not viva_questions:
        return
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Question"
    hdr[2].text = "Expected Answer"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = COLOR_ACCENT
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'E8F0FE')
        cell._tc.get_or_add_tcPr().append(shading)

    for i, qa in enumerate(viva_questions, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = str(qa.get('question', ''))
        row[2].text = str(qa.get('answer', ''))
        for r in row:
            for run in r.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"
    doc.add_paragraph()


def generate_report(project_data):
    doc = Document()

    # ── Page Margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Cover Page ────────────────────────────────────────────────────────────
    # Institution header
    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = inst.add_run("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = COLOR_DARK
    r.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Final Year Project Report")
    rs.font.size = Pt(11)
    rs.font.color.rgb = COLOR_MUTED
    rs.font.name = "Calibri"

    doc.add_paragraph()
    _add_horizontal_rule(doc)
    doc.add_paragraph()

    # Project Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(project_data.get('title', 'Project Report'))
    tr.font.size = Pt(26)
    tr.font.bold = True
    tr.font.color.rgb = COLOR_ACCENT
    tr.font.name = "Calibri"

    doc.add_paragraph()

    # Domain & Difficulty badges
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Domain: {project_data.get('domain', 'Engineering')}    |    "
        f"Difficulty: {project_data.get('difficulty', 'Intermediate')}    |    "
        f"Year: {datetime.now().year}"
    )
    mr.font.size = Pt(12)
    mr.font.color.rgb = COLOR_LIGHT
    mr.font.name = "Calibri"

    doc.add_paragraph()
    _add_horizontal_rule(doc)
    doc.add_paragraph()

    powered = doc.add_paragraph()
    powered.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr2 = powered.add_run("Generated by AxionX AI Platform")
    pr2.font.size = Pt(10)
    pr2.font.color.rgb = COLOR_MUTED
    pr2.font.italic = True
    pr2.font.name = "Calibri"

    doc.add_page_break()

    # ── Table of Contents (Manual) ────────────────────────────────────────────
    h = doc.add_heading("Table of Contents", level=1)
    _style_heading(h, 1, COLOR_ACCENT)
    _add_horizontal_rule(doc)

    toc_items = [
        ("1.", "Abstract"),
        ("2.", "Problem Statement"),
        ("3.", "Literature Survey"),
        ("4.", "Key Features"),
        ("5.", "System Architecture"),
        ("6.", "Database Design"),
        ("7.", "Data Flow & Logic"),
        ("8.", "Security Measures"),
        ("9.", "Technology Stack"),
        ("10.", "Methodology"),
        ("11.", "Viva Questions"),
        ("12.", "Conclusion"),
    ]
    for num, item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        rn = p.add_run(f"  {num}  {item}")
        rn.font.size = Pt(11)
        rn.font.name = "Calibri"
        rn.font.color.rgb = COLOR_LIGHT

    doc.add_page_break()

    # ── Chapter 1: Abstract ───────────────────────────────────────────────────
    _add_chapter(doc, 1, "Abstract")
    _add_body(doc, project_data.get('abstract', 'No abstract available.'))
    doc.add_paragraph()

    # ── Chapter 2: Problem Statement ─────────────────────────────────────────
    _add_chapter(doc, 2, "Problem Statement")
    _add_body(doc, project_data.get('problem_statement', ''))
    doc.add_paragraph()

    # ── Chapter 3: Literature Survey ──────────────────────────────────────────
    lit = project_data.get('literature_survey', '')
    if lit:
        _add_chapter(doc, 3, "Literature Survey")
        _add_body(doc, lit)
        doc.add_paragraph()

    # ── Chapter 4: Key Features ───────────────────────────────────────────────
    features = project_data.get('features', [])
    if features:
        _add_chapter(doc, 4, "Key Features")
        for feature in features:
            _add_bullet(doc, feature)
        doc.add_paragraph()

    # ── Chapter 5: System Architecture ───────────────────────────────────────
    _add_chapter(doc, 5, "System Architecture")
    _add_body(doc, project_data.get('architecture_description', ''))
    doc.add_paragraph()

    # ── Chapter 6: Database Design ────────────────────────────────────────────
    db_design = project_data.get('database_design', '')
    if db_design:
        _add_chapter(doc, 6, "Database Design")
        _add_body(doc, db_design)
        doc.add_paragraph()

    # ── Chapter 7: Data Flow & Logic ─────────────────────────────────────────
    logic = project_data.get('logic_flow', '')
    if logic:
        _add_chapter(doc, 7, "Data Flow & Logic")
        _add_body(doc, logic)
        doc.add_paragraph()

    # ── Chapter 8: Security Measures ─────────────────────────────────────────
    security = project_data.get('security_measures', '')
    if security:
        _add_chapter(doc, 8, "Security Measures")
        _add_body(doc, security)
        doc.add_paragraph()

    # ── Chapter 9: Technology Stack ───────────────────────────────────────────
    tech = project_data.get('tech_stack_details', {})
    _add_chapter(doc, 9, "Technology Stack")
    if isinstance(tech, dict) and tech:
        _add_key_value_table(doc, tech)
    else:
        _add_body(doc, str(tech) if tech else "Refer to project documentation.")
    doc.add_paragraph()

    # ── Chapter 10: Methodology ───────────────────────────────────────────────
    methodology = project_data.get('methodology', '')
    if methodology:
        _add_chapter(doc, 10, "Methodology")
        steps = [s.strip() for s in methodology.split('\n') if s.strip()]
        if len(steps) > 1:
            for step in steps:
                _add_bullet(doc, step)
        else:
            _add_body(doc, methodology)
        doc.add_paragraph()

    # ── Chapter 11: Viva Questions ────────────────────────────────────────────
    viva_qs = project_data.get('viva_questions', [])
    if viva_qs:
        doc.add_page_break()
        _add_chapter(doc, 11, "Viva Questions & Answers")
        _add_viva_table(doc, viva_qs)
        doc.add_paragraph()

    # ── Chapter 12: Conclusion ────────────────────────────────────────────────
    _add_chapter(doc, 12, "Conclusion")
    conclusion_text = (
        f"This project, \"{project_data.get('title', 'the system')}\", successfully demonstrates "
        f"a practical and industry-ready implementation in the domain of {project_data.get('domain', 'Computer Science')}. "
        f"The system was developed with a structured approach covering requirements analysis, architectural planning, "
        f"full implementation, and thorough documentation. "
        f"The solution addresses real-world challenges and is designed to be scalable, secure, and maintainable. "
        f"Future scope includes cloud deployment, mobile integration, and AI-powered enhancements."
    )
    _add_body(doc, conclusion_text)

    doc.add_paragraph()
    _add_horizontal_rule(doc)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("Report generated by AxionX AI Platform | Confidential")
    fr.font.size = Pt(9)
    fr.font.italic = True
    fr.font.color.rgb = COLOR_MUTED
    fr.font.name = "Calibri"

    # ── Save to buffer ────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
